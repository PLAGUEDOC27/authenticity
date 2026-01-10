import os
from flask import Blueprint, request, jsonify, current_app
from flask_jwt_extended import jwt_required, get_jwt_identity
from werkzeug.utils import secure_filename
from flask import render_template

from extensions import db
from models.document import Document

# For text extraction
import docx
from pdfminer.high_level import extract_text
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

document_bp = Blueprint("documents", __name__)

ALLOWED_EXT = {"pdf", "docx", "txt"}

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXT

def extract_text_from_file(filepath, ext):
    if ext == "pdf":
        text = extract_text(filepath)
    elif ext == "docx":
        doc = docx.Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
    elif ext == "txt":
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()
    else:
        text = ""
    return text.strip()

@document_bp.route("/documents/upload", methods=["POST"])
@jwt_required()
def upload_document():
    if "file" not in request.files:
        return jsonify({"msg": "No file part"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"msg": "No selected file"}), 400
    if not allowed_file(file.filename):
        return jsonify({"msg": "Unsupported file type"}), 400

    filename = secure_filename(file.filename)
    upload_folder = os.path.join(current_app.root_path, "uploads")
    os.makedirs(upload_folder, exist_ok=True)
    path = os.path.join(upload_folder, filename)
    file.save(path)

    ext = filename.rsplit(".", 1)[1].lower()
    user_id = get_jwt_identity()

    # Extract text
    original_text = extract_text_from_file(path, ext)
    text_length = len(original_text)

    # Create Document entry
    doc = Document(
        user_id=user_id,
        filename=filename,
        file_type=ext,
        original_text=original_text
    )
    db.session.add(doc)
    db.session.commit()

    # --- Plagiarism check (TF-IDF + Cosine Similarity) ---
    # Get all other documents
    other_docs = Document.query.filter(Document.id != doc.id).all()
    max_similarity = 0
    if other_docs:
        corpus = [doc.original_text] + [d.original_text for d in other_docs]
        vectorizer = TfidfVectorizer().fit_transform(corpus)
        similarity_matrix = cosine_similarity(vectorizer[0:1], vectorizer[1:])
        max_similarity = float(similarity_matrix.max())  # between 0-1

    # --- AI-text detection placeholder ---
    # For now, just a random placeholder (later we can use a real ML model)
    ai_generated_prob = 0.0  # 0-1 probability

    return jsonify({
        "document_id": doc.id,
        "msg": "File uploaded and text extracted",
        "text_length": text_length,
        "max_similarity": max_similarity,
        "ai_generated_prob": ai_generated_prob
    }), 201
    
@document_bp.route("/dashboard", methods=["GET"])
def dashboard():
    documents = Document.query.all()

    # Prepare documents for display
    docs_for_display = []
    for d in documents:
        docs_for_display.append({
            "id": d.id,
            "filename": d.filename,
            "user_id": d.user_id,
            "text_length": len(d.original_text) if d.original_text else 0,
            "max_similarity": getattr(d, "max_similarity", 0.0),  # placeholder
            "ai_generated_prob": getattr(d, "ai_generated_prob", 0.0)  # placeholder
        })

    return render_template("dashboard.html", documents=docs_for_display)